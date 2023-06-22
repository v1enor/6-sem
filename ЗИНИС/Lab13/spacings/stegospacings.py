from docx import Document
from bidict import bidict
from utils import *

spacings = (1.0, 1.1, 1.2, 1.3)
spacing_codes = bidict({
    1.0: '00',
    1.1: '01',
    1.2: '10',
    1.3: '11'
})


def encode(message, container_path, output_path):
    doc = Document(container_path)
    bits = to_bits(message)
    
    if len(doc.paragraphs) < (int(len(bits) / 2)) :
        raise Exception('No enough paragraphs in container!')
    
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.08
    
    i = 0
    while bits != '':
        doc.paragraphs[i].paragraph_format.line_spacing = spacing_codes.inverse[bits[:2]]
        bits = bits[2:]
        i += 1
    
    doc.save(output_path)


def decode(path):
    doc = Document(path)

    bits = ''
    for para in doc.paragraphs:
        if para.paragraph_format.line_spacing in spacings:
            bits += spacing_codes[para.paragraph_format.line_spacing]
    
    return from_bits(bits)
